"""
Generate professionally styled Word document for ChronosDB AI Layer Documentation.
Uses python-docx for document creation with custom styling.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Color Palette ──────────────────────────────────────────────────
PRIMARY      = RGBColor(0x1A, 0x56, 0xDB)   # Deep blue
ACCENT       = RGBColor(0x0E, 0x9F, 0x6E)   # Teal green
DARK         = RGBColor(0x1F, 0x29, 0x37)   # Near-black
MEDIUM       = RGBColor(0x4B, 0x55, 0x63)   # Dark gray
LIGHT_TEXT   = RGBColor(0x6B, 0x72, 0x80)   # Gray
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG      = RGBColor(0xF3, 0xF4, 0xF6)   # Light gray for code
TABLE_HEADER = RGBColor(0x1A, 0x56, 0xDB)   # Blue header
TABLE_ALT    = RGBColor(0xF0, 0xF4, 0xFF)   # Light blue alternating

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" '
            f'w:sz="{val["sz"]}" w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def style_document(doc):
    """Configure document-level styles."""
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = DARK
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = DARK
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = ACCENT
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    # Heading 4
    h4 = doc.styles['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.color.rgb = MEDIUM
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, '1A56DB')

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4FF')

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    # Table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # spacing after table
    return table


def add_code_block(doc, code, language=''):
    """Add a styled code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK

    # Add shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F3F4F6"/>')
    pPr.append(shading)

    return p


def add_bullet(doc, text, bold_prefix='', level=0):
    """Add a bullet point with optional bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    else:
        p.runs[0].font.size = Pt(11) if p.runs else None
        if not p.runs:
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    return p


def add_horizontal_rule(doc):
    """Add a horizontal line separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_info_box(doc, text, box_color='E8F4FD', border_color='1A56DB'):
    """Add an info/callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run.italic = True

    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{box_color}"/>')
    pPr.append(shading)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="4" w:color="{border_color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def build_document():
    doc = Document()
    style_document(doc)

    # ═══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ChronosDB')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('AI Layer Technical Documentation')
    run.font.size = Pt(20)
    run.font.color.rgb = MEDIUM
    run.font.name = 'Calibri'

    doc.add_paragraph()

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run('\u2501' * 40)
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(14)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Self-Learning Execution Engine  |  Immune System  |  Temporal Index Manager')
    run.font.size = Pt(12)
    run.font.color.rgb = LIGHT_TEXT
    run.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Version 1.0  \u2022  February 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = LIGHT_TEXT
    run.font.name = 'Calibri'

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run('C++20  \u2022  CMake + Ninja  \u2022  Windows / Linux')
    run.font.size = Pt(11)
    run.font.color.rgb = LIGHT_TEXT
    run.font.name = 'Calibri'

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Project Overview'),
        ('2.', 'System Architecture'),
        ('3.', 'AI Layer Overview'),
        ('4.', 'Phase 0: Shared Foundation'),
        ('5.', 'Phase 1: Self-Learning Execution Engine'),
        ('6.', 'Phase 2: Immune System'),
        ('7.', 'Phase 3: Intelligent Temporal Index Manager'),
        ('8.', 'Integration Architecture'),
        ('9.', 'SQL Commands'),
        ('10.', 'UML Diagrams'),
        ('11.', 'File Inventory'),
        ('12.', 'Algorithm Reference'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f'{num}  ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = PRIMARY
        run.font.name = 'Calibri'
        run = p.add_run(title_text)
        run.font.size = Pt(12)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('1. Project Overview', level=1)

    p = doc.add_paragraph()
    run = p.add_run('ChronosDB')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(' is a high-performance, multi-protocol database management system written in C++20. '
                     'It is designed from the ground up to support temporal data operations with built-in '
                     'time-travel capabilities.')
    run.font.size = Pt(11)

    doc.add_heading('Core Features', level=3)
    features = [
        ('Role-Based Access Control (RBAC): ', '5 roles \u2014 SUPERADMIN, ADMIN, USER, READONLY, DENIED'),
        ('Persistent Storage: ', 'Crash recovery via Write-Ahead Logging (WAL)'),
        ('Time Travel: ', 'Query data at any historical point using SELECT ... AS OF'),
        ('CQL (Chronos Query Language): ', 'SQL with Franco-Arabic keyword alternatives'),
        ('B+ Tree Indexes: ', 'Fast ordered lookups with page-level locking'),
        ('Buffer Pool Management: ', 'Adaptive partitioned strategy with multiple eviction policies'),
        ('Transaction Support: ', 'Auto-commit and explicit BEGIN/COMMIT/ROLLBACK'),
    ]
    for prefix, desc_text in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run = p.add_run(desc_text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    doc.add_heading('Execution Flow', level=3)
    p = doc.add_paragraph('The query execution pipeline follows a clean, layered architecture:')
    add_code_block(doc,
        'Client (Shell/Network)\n'
        '    \u2502\n'
        '    \u25bc\n'
        'ConnectionHandler (Multi-protocol: TEXT, JSON, BINARY)\n'
        '    \u2502\n'
        '    \u25bc\n'
        'Parser (Lexer \u2192 Token Stream \u2192 AST)\n'
        '    \u2502\n'
        '    \u25bc\n'
        'ExecutionEngine (Dispatch Map Pattern)\n'
        '    \u2502\n'
        '    \u251c\u2500\u2500\u2500 DDLExecutor     (CREATE, DROP, ALTER)\n'
        '    \u251c\u2500\u2500\u2500 DMLExecutor     (INSERT, SELECT, UPDATE, DELETE)\n'
        '    \u251c\u2500\u2500\u2500 SystemExecutor  (SHOW, WHOAMI, STATUS)\n'
        '    \u251c\u2500\u2500\u2500 UserExecutor    (CREATE USER, ALTER USER)\n'
        '    \u251c\u2500\u2500\u2500 DatabaseExecutor(CREATE DATABASE, USE)\n'
        '    \u2514\u2500\u2500\u2500 TransactionExecutor (BEGIN, COMMIT, ROLLBACK)\n'
        '    \u2502\n'
        '    \u25bc\n'
        'Storage Layer (BufferPoolManager \u2192 DiskManager \u2192 Pages)\n'
        '    \u2502\n'
        '    \u25bc\n'
        'Recovery Layer (LogManager \u2192 WAL \u2192 CheckpointManager)')

    doc.add_heading('Key Technologies', level=3)
    add_styled_table(doc,
        ['Component', 'Technology'],
        [
            ['Language', 'C++20'],
            ['Build System', 'CMake + Ninja'],
            ['Page Size', '4 KB'],
            ['Default Buffer Pool', '65,536 pages (256 MB)'],
            ['Network Port', '2501'],
            ['Protocol', 'chronos://user:pass@host:port/db'],
        ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('2. System Architecture', level=1)

    doc.add_heading('Design Patterns', level=2)
    p = doc.add_paragraph('ChronosDB follows SOLID principles and uses well-established design patterns '
                           'throughout the codebase:')

    add_styled_table(doc,
        ['Pattern', 'Where Used', 'Purpose'],
        [
            ['Factory', 'ExecutorFactory', 'Creates executors based on statement type'],
            ['Dispatch Map', 'ExecutionEngine', 'Routes statements to handlers without switch/if-else chains'],
            ['Observer', 'DMLObserverRegistry', 'Decouples AI layer from execution engine'],
            ['Strategy', 'IQueryOptimizer', 'Swappable scan strategy selection at runtime'],
            ['Singleton', 'AIManager, MetricsStore, AIScheduler', 'Global AI coordination with single instance'],
            ['RAII', 'Throughout codebase', 'Resource management via smart pointers'],
            ['Iterator', 'TableHeap::Iterator', 'Tuple-at-a-time row processing'],
        ])

    doc.add_heading('Thread Safety Model', level=2)
    safety_items = [
        ('Global Reader-Writer Lock: ', 'ExecutionEngine::global_lock_ protects concurrent DML (shared lock) vs. exclusive recovery operations (exclusive lock).'),
        ('Per-component Mutexes: ', 'Each AI subsystem uses std::shared_mutex for read-heavy access patterns, allowing multiple concurrent readers with single writer exclusion.'),
        ('Lock-Free Ring Buffer: ', 'MetricsStore uses std::atomic write index for zero-contention metric recording on the hot path.'),
        ('Atomic Flags: ', 'All AI subsystems use std::atomic<bool> for safe start/stop lifecycle transitions without mutex overhead.'),
    ]
    for prefix, desc_text in safety_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run = p.add_run(desc_text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 3. AI LAYER OVERVIEW
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('3. AI Layer Overview', level=1)

    p = doc.add_paragraph('The AI layer consists of three integrated subsystems that share infrastructure and '
                           'feed each other data. They are not bolted-on features \u2014 they are embedded in the '
                           'database kernel and operate on the critical path of every DML operation.')

    add_info_box(doc, 'Key Principle: The AI layer uses the Observer pattern to hook into every DML operation. '
                      'This ensures zero coupling between the core execution engine and the AI subsystems. '
                      'AI components register themselves as observers and are notified before and after every operation.')

    doc.add_heading('Three AI Subsystems', level=2)

    # System 1
    doc.add_heading('Self-Learning Execution Engine (Phase 1)', level=3)
    p = doc.add_paragraph('Uses a UCB1 Multi-Armed Bandit algorithm to learn which scan strategy '
                           '(Sequential Scan vs. Index Scan) is optimal for each query pattern. '
                           'No manual tuning required \u2014 the system learns from real execution feedback.')

    # System 2
    doc.add_heading('Immune System (Phase 2)', level=3)
    p = doc.add_paragraph('Provides autonomous anomaly detection and self-healing. Monitors mutation patterns '
                           'across all tables, detects anomalous behavior using Z-score statistical analysis, '
                           'and can automatically block operations or recover data using the Time Travel Engine.')

    # System 3
    doc.add_heading('Intelligent Temporal Index Manager (Phase 3)', level=3)
    p = doc.add_paragraph('Makes time-travel queries faster by intelligently scheduling snapshots. '
                           'Uses DBSCAN clustering to detect temporal hotspots and CUSUM change-point detection '
                           'to identify optimal snapshot placement moments.')

    doc.add_heading('How the Three Systems Interconnect', level=2)
    interconnect = [
        ('MetricsStore', ' is the shared data backbone \u2014 all three systems record events into the same lock-free ring buffer.'),
        ('DMLObserverRegistry', ' broadcasts every INSERT/UPDATE/DELETE/SELECT to both the Learning Engine and Immune System simultaneously.'),
        ('Learning Engine', ' learns optimal scan strategies from SELECT feedback and feeds recommendations back into DMLExecutor.'),
        ('Immune System', ' monitors mutation patterns and can trigger auto-recovery via TimeTravelEngine when anomalies are detected.'),
        ('Temporal Index Manager', ' monitors time-travel query patterns and optimizes snapshot placement for faster recovery operations.'),
    ]
    for i, (prefix, desc_text) in enumerate(interconnect):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{i+1}. ')
        run.bold = True
        run.font.color.rgb = PRIMARY
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(desc_text)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 4. PHASE 0: SHARED FOUNDATION
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('4. Phase 0: Shared Foundation', level=1)

    p = doc.add_paragraph('The shared foundation provides the infrastructure that all three AI subsystems depend on. '
                           'This includes centralized configuration, a shared metrics store, the observer framework, '
                           'a background task scheduler, and the top-level AI coordinator.')

    # 4.1 AIConfig
    doc.add_heading('4.1 AIConfig (ai_config.h)', level=2)
    p = doc.add_paragraph('All AI constants are centralized in a single header file. This eliminates magic numbers '
                           'in implementation files and provides a single point of configuration.')
    add_styled_table(doc,
        ['Constant', 'Value', 'Purpose'],
        [
            ['UCB1_EXPLORATION_CONSTANT', '\u221a2 \u2248 1.414', 'UCB1 exploration vs. exploitation tradeoff'],
            ['MIN_SAMPLES_BEFORE_LEARNING', '30', 'Queries before AI starts recommending strategies'],
            ['MIN_ARM_PULLS', '5', 'Minimum pulls per arm before exploitation'],
            ['ZSCORE_LOW_THRESHOLD', '2.0', 'Anomaly detection: log warning level'],
            ['ZSCORE_MEDIUM_THRESHOLD', '3.0', 'Anomaly detection: block mutations'],
            ['ZSCORE_HIGH_THRESHOLD', '4.0', 'Anomaly detection: auto-recover'],
            ['IMMUNE_CHECK_INTERVAL_MS', '1,000', 'Immune system check frequency (1s)'],
            ['HOTSPOT_CLUSTER_EPSILON_US', '60,000,000', 'DBSCAN: 60-second neighborhood'],
            ['HOTSPOT_CLUSTER_MIN_POINTS', '5', 'DBSCAN: minimum cluster size'],
            ['CUSUM_THRESHOLD_SIGMA_MULT', '4.0', 'CUSUM: change detection threshold'],
            ['METRICS_RING_BUFFER_SIZE', '8,192', 'Maximum metrics in ring buffer'],
            ['ACCESS_PATTERN_WINDOW_SIZE', '1,000', 'Temporal access history size'],
        ])

    # 4.2 MetricsStore
    doc.add_heading('4.2 MetricsStore', level=2)
    p = doc.add_paragraph('Thread-safe ring buffer for operation metrics, shared by all 3 AI subsystems. '
                           'Uses a lock-free write mechanism via atomic index for zero-contention on the hot path.')

    doc.add_heading('Class: MetricsStore (Singleton)', level=4)
    add_code_block(doc,
        'class MetricsStore {\n'
        'public:\n'
        '    static MetricsStore& Instance();\n'
        '    void Record(const MetricEvent& event);          // Lock-free write\n'
        '    vector<MetricEvent> Query(MetricType type, uint64_t since_us) const;\n'
        '    size_t CountEvents(MetricType type, uint64_t since_us) const;\n'
        '    uint64_t GetMutationCount(const string& table, uint64_t since_us) const;\n'
        '    size_t GetTotalRecorded() const;\n'
        '    void Reset();\n'
        '};', 'cpp')

    doc.add_heading('MetricEvent Structure', level=4)
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['type', 'MetricType', 'INSERT, UPDATE, DELETE, SELECT, ANOMALY_DETECTED, etc.'],
            ['timestamp_us', 'uint64_t', 'Microsecond timestamp of the event'],
            ['duration_us', 'uint64_t', 'Operation duration in microseconds'],
            ['user', 'string', 'User who executed the operation'],
            ['table_name', 'string', 'Target table name'],
            ['db_name', 'string', 'Target database name'],
            ['rows_affected', 'uint32_t', 'Number of rows changed'],
            ['scan_strategy', 'uint8_t', '0 = SeqScan, 1 = IndexScan'],
            ['target_timestamp', 'uint64_t', 'For time-travel queries: the AS OF timestamp'],
        ])

    add_info_box(doc, 'Implementation Detail: The ring buffer uses a fixed-size array of 8,192 entries with an '
                      'atomic write index. Writes never block. Oldest entries are silently overwritten when the '
                      'buffer wraps around. Reads use std::shared_mutex for bulk query operations only.')

    # 4.3 DMLObserverRegistry
    doc.add_heading('4.3 DMLObserverRegistry', level=2)
    p = doc.add_paragraph('The Observer pattern implementation that decouples AI from the execution engine. '
                           'Both the Learning Engine and Immune System implement the IDMLObserver interface '
                           'and register themselves with this registry.')

    doc.add_heading('Interface: IDMLObserver', level=4)
    add_code_block(doc,
        'class IDMLObserver {\n'
        '    virtual bool OnBeforeDML(const DMLEvent& event) { return true; }  // false = block\n'
        '    virtual void OnAfterDML(const DMLEvent& event) {}\n'
        '};', 'cpp')

    doc.add_heading('Class: DMLObserverRegistry (Singleton)', level=4)
    add_code_block(doc,
        'class DMLObserverRegistry {\n'
        '    void Register(IDMLObserver* observer);\n'
        '    void Unregister(IDMLObserver* observer);\n'
        '    bool NotifyBefore(const DMLEvent& event);  // Returns false if ANY observer blocks\n'
        '    void NotifyAfter(const DMLEvent& event);\n'
        '};', 'cpp')

    doc.add_heading('DMLEvent Structure', level=4)
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['operation', 'DMLOperation', 'INSERT, UPDATE, DELETE_OP, SELECT'],
            ['table_name', 'string', 'Target table'],
            ['db_name', 'string', 'Current database'],
            ['user', 'string', 'Current user'],
            ['session_id', 'uint32_t', 'Session identifier'],
            ['rows_affected', 'uint32_t', 'Rows changed (filled after operation)'],
            ['start_time_us', 'uint64_t', 'Operation start time'],
            ['duration_us', 'uint64_t', 'Elapsed time (filled after operation)'],
            ['used_index_scan', 'bool', 'Whether index scan was used (SELECT only)'],
            ['where_clause_count', 'size_t', 'Number of WHERE conditions (SELECT only)'],
            ['has_order_by', 'bool', 'ORDER BY present (SELECT only)'],
            ['has_limit', 'bool', 'LIMIT present (SELECT only)'],
            ['result_row_count', 'int32_t', 'Rows returned (SELECT only)'],
        ])

    # 4.4 AIScheduler
    doc.add_heading('4.4 AIScheduler', level=2)
    p = doc.add_paragraph('Background task manager for periodic AI analysis. Uses a dedicated scheduler thread '
                           'with fine-grained sleep ticks for responsive shutdown.')
    add_code_block(doc,
        'class AIScheduler {\n'
        '    TaskId SchedulePeriodic(const string& name, uint32_t interval_ms,\n'
        '                           function<void()> task);\n'
        '    TaskId ScheduleOnce(const string& name, uint32_t delay_ms,\n'
        '                       function<void()> task);\n'
        '    void Cancel(TaskId id);\n'
        '    size_t GetActiveTaskCount() const;\n'
        '};', 'cpp')
    add_info_box(doc, 'The scheduler thread sleeps in 10ms ticks and checks a running_ atomic flag each tick. '
                      'This ensures shutdown completes within 10ms even when tasks have long intervals.')

    # 4.5 AIManager
    doc.add_heading('4.5 AIManager', level=2)
    p = doc.add_paragraph('Top-level singleton coordinator for all AI subsystems. Provides the single entry point '
                           'for AI lifecycle management and status queries.')
    add_code_block(doc,
        'class AIManager {\n'
        '    static AIManager& Instance();\n'
        '    void Initialize(Catalog*, IBufferManager*, LogManager*, CheckpointManager*);\n'
        '    void Shutdown();\n'
        '    bool IsInitialized() const;\n'
        '    LearningEngine* GetLearningEngine();\n'
        '    ImmuneSystem* GetImmuneSystem();\n'
        '    TemporalIndexManager* GetTemporalIndexManager();\n'
        '    AIStatus GetStatus() const;\n'
        '};', 'cpp')

    doc.add_heading('Lifecycle', level=4)
    lifecycle = [
        'Initialize() is called in ExecutionEngine constructor',
        'Creates all 3 subsystems in order: LearningEngine \u2192 ImmuneSystem \u2192 TemporalIndexManager',
        'Registers LearningEngine and ImmuneSystem as DML observers',
        'Starts all 3 subsystems (schedules periodic background tasks)',
        'Shutdown() stops all subsystems, unregisters observers, destroys in reverse order',
    ]
    for i, item in enumerate(lifecycle):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{i+1}. ')
        run.bold = True
        run.font.color.rgb = PRIMARY
        run = p.add_run(item)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 5. PHASE 1: SELF-LEARNING EXECUTION ENGINE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('5. Phase 1: Self-Learning Execution Engine', level=1)

    p = doc.add_paragraph('The Learning Engine uses a UCB1 Multi-Armed Bandit algorithm to learn which scan strategy '
                           'is optimal for each query pattern. It treats Sequential Scan and Index Scan as two '
                           '"arms" of a slot machine and uses real execution feedback to learn which performs better.')

    # 5.1 QueryFeatureExtractor
    doc.add_heading('5.1 QueryFeatureExtractor', level=2)
    p = doc.add_paragraph('Extracts an 8-dimensional feature vector from a SelectStatement, providing '
                           'the bandit with contextual information about each query.')

    add_styled_table(doc,
        ['Feature', 'Type', 'Description'],
        [
            ['table_row_count_log', 'double', 'log2(estimated row count)'],
            ['where_clause_count', 'double', 'Number of WHERE conditions'],
            ['has_equality_predicate', 'double', '1.0 if any WHERE uses = operator'],
            ['has_index_available', 'double', '1.0 if index exists on first WHERE column'],
            ['selectivity_estimate', 'double', 'Estimated fraction of rows matching (0.1 for =, 0.33 for range)'],
            ['column_count', 'double', 'Number of selected columns'],
            ['has_order_by', 'double', '1.0 if ORDER BY present'],
            ['has_limit', 'double', '1.0 if LIMIT present'],
        ])

    # 5.2 UCB1Bandit
    doc.add_heading('5.2 UCB1 Bandit Algorithm', level=2)

    doc.add_heading('Selection Formula', level=4)
    add_code_block(doc,
        'UCB1 Selection:\n'
        '    score(a) = Q(a) + c \u00b7 \u221a(ln(N) / N_a)\n'
        '\n'
        '    Where:\n'
        '        Q(a) = average reward for arm a = total_reward(a) / N_a\n'
        '        c    = exploration constant = \u221a2 \u2248 1.414\n'
        '        N    = total pulls across all arms\n'
        '        N_a  = pulls for arm a\n'
        '\n'
        '    Select: argmax_a [ score(a) ]')

    doc.add_heading('Reward Function', level=4)
    add_code_block(doc,
        'reward(time_ms) = 1.0 / (1.0 + time_ms / 100.0)\n'
        '\n'
        'Examples:\n'
        '    10ms query   \u2192 reward = 0.91 (excellent)\n'
        '    100ms query  \u2192 reward = 0.50 (average)\n'
        '    1000ms query \u2192 reward = 0.09 (poor)\n'
        '\n'
        'Properties: Always in (0, 1], smooth, bounded, differentiable')

    doc.add_heading('Per-Table Contextual Learning', level=4)
    ctx_items = [
        'Global statistics track overall arm performance across all tables',
        'Per-table statistics override globals after 10+ pulls per arm for that specific table',
        'Each table independently learns its own optimal scan strategy',
        'Fixed-point arithmetic (total_reward_x10000) avoids floating-point atomic issues',
    ]
    for item in ctx_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    doc.add_heading('Exploration Strategy', level=4)
    explore_items = [
        'First 30 queries: no recommendation (use database\u2019s existing heuristic behavior)',
        'Per-arm minimum 5 pulls: ensures each strategy gets tested before exploitation begins',
        'After sufficient data: UCB1 formula balances exploitation of the best arm with exploration of uncertain arms',
    ]
    for item in explore_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # 5.3 LearningEngine
    doc.add_heading('5.3 LearningEngine Orchestrator', level=2)
    p = doc.add_paragraph('The orchestrator that connects the bandit algorithm to the DML execution path. '
                           'It implements both IDMLObserver (to receive execution feedback) and IQueryOptimizer '
                           '(to provide scan strategy recommendations).')

    add_code_block(doc,
        'class LearningEngine : public IDMLObserver, public IQueryOptimizer {\n'
        '    // IDMLObserver \u2014 called after every SELECT\n'
        '    void OnAfterDML(const DMLEvent& event) override;\n'
        '\n'
        '    // IQueryOptimizer \u2014 consulted before scan strategy decision\n'
        '    bool RecommendScanStrategy(const SelectStatement* stmt,\n'
        '                               const string& table_name,\n'
        '                               ScanStrategy& out_strategy) override;\n'
        '\n'
        '    string GetSummary() const;\n'
        '    vector<ArmStats> GetArmStats() const;\n'
        '    uint64_t GetTotalQueriesObserved() const;\n'
        '};', 'cpp')

    doc.add_heading('Execution Flow', level=4)
    add_code_block(doc,
        'DMLExecutor::Select()\n'
        '    \u2502\n'
        '    \u251c\u2500\u2500\u2500 LearningEngine::RecommendScanStrategy()\n'
        '    \u2502         \u2502\n'
        '    \u2502         \u251c\u2500\u2500\u2500 QueryFeatureExtractor::Extract()\n'
        '    \u2502         \u2514\u2500\u2500\u2500 UCB1Bandit::SelectStrategy(table_name)\n'
        '    \u2502                 \u2514\u2500\u2500\u2500 Returns INDEX_SCAN or SEQUENTIAL_SCAN\n'
        '    \u2502\n'
        '    \u251c\u2500\u2500\u2500 Execute query with recommended strategy\n'
        '    \u2502\n'
        '    \u2514\u2500\u2500\u2500 DMLObserverRegistry::NotifyAfter(event)\n'
        '              \u2502\n'
        '              \u2514\u2500\u2500\u2500 LearningEngine::OnAfterDML()\n'
        '                        \u251c\u2500\u2500\u2500 UCB1Bandit::RecordOutcome(strategy, time)\n'
        '                        \u2514\u2500\u2500\u2500 MetricsStore::Record(metric)')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 6. PHASE 2: IMMUNE SYSTEM
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('6. Phase 2: Immune System', level=1)

    p = doc.add_paragraph('The Immune System provides autonomous anomaly detection and self-healing for the database. '
                           'Like a biological immune system, it monitors the "health" of the database by tracking '
                           'mutation patterns, detecting statistical anomalies, profiling user behavior, and taking '
                           'graduated defensive actions up to automatic data recovery.')

    # 6.1 MutationMonitor
    doc.add_heading('6.1 MutationMonitor', level=2)
    p = doc.add_paragraph('Per-table rolling window of mutation events for rate calculation. Maintains a sliding '
                           'window of mutation timestamps and row counts per table.')

    add_code_block(doc,
        'class MutationMonitor {\n'
        '    void RecordMutation(const string& table_name, uint32_t row_count);\n'
        '    double GetMutationRate(const string& table_name) const;      // mutations/sec\n'
        '    vector<double> GetHistoricalRates(const string& table_name) const;\n'
        '    vector<string> GetMonitoredTables() const;\n'
        '};', 'cpp')

    impl_items = [
        ('Rolling window: ', '10-minute sliding window using std::deque<MutationEntry>'),
        ('Rate calculation: ', 'Counts mutations in the last 5-second window'),
        ('Historical rates: ', 'Divides history into 1-second intervals for Z-score computation'),
        ('Auto-pruning: ', 'Entries older than the rolling window are automatically removed'),
    ]
    for prefix, desc_text in impl_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run = p.add_run(desc_text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # 6.2 UserBehaviorProfiler
    doc.add_heading('6.2 UserBehaviorProfiler', level=2)
    p = doc.add_paragraph('Per-user behavioral baselines for detecting compromised accounts or unauthorized access patterns.')

    add_code_block(doc,
        'class UserBehaviorProfiler {\n'
        '    void RecordActivity(const string& user, const string& table_name,\n'
        '                        bool is_mutation);\n'
        '    double GetDeviationScore(const string& user) const;\n'
        '    vector<UserProfile> GetAllProfiles() const;\n'
        '};', 'cpp')

    doc.add_heading('Deviation Score Formula', level=4)
    add_code_block(doc,
        'deviation = 0.6 \u00d7 mutation_rate_deviation + 0.4 \u00d7 table_access_deviation\n'
        '\n'
        'mutation_rate_deviation:\n'
        '    current_rate = mutations in last 5 seconds\n'
        '    history = rates over last 100 intervals\n'
        '    z = (current_rate - mean(history)) / stddev(history)\n'
        '\n'
        'table_access_deviation:\n'
        '    unique_tables_now = tables accessed in current window\n'
        '    if unique_tables_now > historical_max \u00d7 2 \u2192 high deviation')

    # 6.3 AnomalyDetector
    doc.add_heading('6.3 AnomalyDetector', level=2)
    p = doc.add_paragraph('Z-score computation on mutation rates with severity classification. '
                           'This is the core statistical engine of the Immune System.')

    doc.add_heading('Z-Score Algorithm', level=4)
    add_code_block(doc,
        'z = (x - \u03bc) / \u03c3\n'
        '\n'
        'Where:\n'
        '    x = current mutation rate (mutations/sec in last check interval)\n'
        '    \u03bc = mean of last 100 intervals\n'
        '    \u03c3 = standard deviation of last 100 intervals')

    doc.add_heading('Severity Classification', level=4)
    add_styled_table(doc,
        ['Z-Score Range', 'Severity', 'Statistical Meaning', 'Action'],
        [
            ['z < 2.0', 'NONE', 'Within normal variation (95.4%)', 'No action'],
            ['z \u2265 2.0', 'LOW', 'Outside 2\u03c3 (4.6% false positive)', 'Log warning'],
            ['z \u2265 3.0', 'MEDIUM', 'Outside 3\u03c3 (0.3% false positive)', 'Block table mutations'],
            ['z \u2265 4.0', 'HIGH', 'Outside 4\u03c3 (0.006% false positive)', 'Auto-recover via Time Travel'],
        ])

    doc.add_heading('AnomalyReport Structure', level=4)
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['table_name', 'string', 'Affected table'],
            ['user', 'string', 'Most active user during anomaly'],
            ['severity', 'AnomalySeverity', 'NONE / LOW / MEDIUM / HIGH'],
            ['z_score', 'double', 'Computed Z-score value'],
            ['current_rate', 'double', 'Current mutations/sec'],
            ['mean_rate', 'double', 'Historical mean mutations/sec'],
            ['std_dev', 'double', 'Historical standard deviation'],
            ['timestamp_us', 'uint64_t', 'When the anomaly was detected'],
            ['description', 'string', 'Human-readable description'],
        ])

    # 6.4 ResponseEngine
    doc.add_heading('6.4 ResponseEngine', level=2)
    p = doc.add_paragraph('Executes graduated responses based on anomaly severity. Implements a defense-in-depth '
                           'strategy with escalating countermeasures.')

    doc.add_heading('Response Ladder', level=4)
    add_styled_table(doc,
        ['Severity', 'Action', 'Description'],
        [
            ['LOW', 'LOG_WARN', 'Log warning with Z-score and mutation rates for monitoring'],
            ['MEDIUM', 'BLOCK', 'Block all mutations on the affected table and user'],
            ['HIGH', 'AUTO-RECOVER', 'Block + trigger TimeTravelEngine::RecoverTo() to 60s before anomaly'],
        ])

    doc.add_heading('Auto-Recovery Flow (HIGH Severity)', level=4)
    recovery = [
        'Block the table immediately to prevent further damage',
        'Calculate recovery timestamp: anomaly_time \u2212 60 seconds',
        'Call TimeTravelEngine::RecoverTo(timestamp, db_name)',
        'If recovery succeeds, unblock the table',
        'If recovery fails, keep table blocked and log error for manual intervention',
    ]
    for i, item in enumerate(recovery):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{i+1}. ')
        run.bold = True
        run.font.color.rgb = PRIMARY
        run = p.add_run(item)
        run.font.size = Pt(11)

    # 6.5 ImmuneSystem
    doc.add_heading('6.5 ImmuneSystem Orchestrator', level=2)
    p = doc.add_paragraph('The main orchestrator that ties together all immune components. Implements IDMLObserver '
                           'to intercept every DML operation.')

    add_code_block(doc,
        'class ImmuneSystem : public IDMLObserver {\n'
        '    bool OnBeforeDML(const DMLEvent& event) override;  // Can block!\n'
        '    void OnAfterDML(const DMLEvent& event) override;   // Records mutations\n'
        '    void PeriodicAnalysis();  // Runs every 1 second\n'
        '    string GetSummary() const;\n'
        '    vector<AnomalyReport> GetRecentAnomalies(size_t max) const;\n'
        '};', 'cpp')

    doc.add_heading('Periodic Analysis Pipeline', level=4)
    add_code_block(doc,
        'Every 1 second:\n'
        '    1. AnomalyDetector::Analyze(mutation_monitor)\n'
        '       \u2192 Returns list of AnomalyReport for each monitored table\n'
        '\n'
        '    2. For each report with severity > NONE:\n'
        '       \u2192 ResponseEngine::Respond(report)\n'
        '       \u2192 Store in anomaly history for SHOW ANOMALIES\n'
        '\n'
        '    3. Record anomaly metrics in MetricsStore')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 7. PHASE 3: INTELLIGENT TEMPORAL INDEX MANAGER
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('7. Phase 3: Intelligent Temporal Index Manager', level=1)

    p = doc.add_paragraph('The Temporal Index Manager makes time-travel queries faster by intelligently scheduling '
                           'snapshots. It tracks which historical timestamps are queried most frequently, detects '
                           'temporal hotspots using DBSCAN clustering, and identifies optimal snapshot placement '
                           'using CUSUM change-point detection.')

    # 7.1 TemporalAccessTracker
    doc.add_heading('7.1 TemporalAccessTracker', level=2)
    p = doc.add_paragraph('Records which timestamps are queried in time-travel operations and provides '
                           'frequency analysis capabilities.')
    add_code_block(doc,
        'class TemporalAccessTracker {\n'
        '    void RecordAccess(const TemporalAccessEvent& event);\n'
        '    vector<FrequencyBucket> GetFrequencyHistogram(uint64_t bucket_width_us) const;\n'
        '    vector<uint64_t> GetHotTimestamps(size_t k) const;\n'
        '    vector<TemporalAccessEvent> GetAllEvents() const;\n'
        '    size_t GetTotalAccessCount() const;\n'
        '};', 'cpp')

    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['queried_timestamp_us', 'uint64_t', 'The historical timestamp being queried (AS OF)'],
            ['query_time_us', 'uint64_t', 'When the time-travel query was executed'],
            ['table_name', 'string', 'Which table was queried'],
            ['db_name', 'string', 'Which database'],
        ])

    # 7.2 HotspotDetector
    doc.add_heading('7.2 HotspotDetector', level=2)
    p = doc.add_paragraph('Detects temporal hotspots using DBSCAN clustering and change points using CUSUM. '
                           'These two algorithms work together to identify both where users are looking '
                           '(hotspots) and when patterns shift (change points).')

    doc.add_heading('DBSCAN Clustering (Simplified 1D)', level=3)
    add_code_block(doc,
        'Input: Sorted timestamps, epsilon (60s), min_points (5)\n'
        '\n'
        'Algorithm:\n'
        '    1. Sort queried timestamps\n'
        '    2. Initialize current_cluster = [first_point]\n'
        '    3. For each subsequent point:\n'
        '       - If distance to previous point \u2264 epsilon (60s):\n'
        '         Add to current cluster\n'
        '       - Else:\n'
        '         If current_cluster.size \u2265 min_points (5):\n'
        '           Save cluster as hotspot\n'
        '         Start new cluster\n'
        '    4. Check last cluster\n'
        '\n'
        'Complexity: O(n) for sorted 1D data (vs. O(n\u00b2) for general DBSCAN)')

    doc.add_heading('TemporalHotspot Structure', level=4)
    add_styled_table(doc,
        ['Field', 'Type', 'Description'],
        [
            ['center_timestamp_us', 'uint64_t', 'Center of the hotspot (mean of cluster)'],
            ['range_start_us', 'uint64_t', 'Earliest timestamp in cluster'],
            ['range_end_us', 'uint64_t', 'Latest timestamp in cluster'],
            ['access_count', 'size_t', 'Number of queries hitting this hotspot'],
            ['density', 'double', 'Queries per second within the hotspot range'],
        ])

    doc.add_heading('CUSUM Change-Point Detection', level=3)
    add_code_block(doc,
        'Input: Time series of mutation rates\n'
        '\n'
        'Parameters:\n'
        '    threshold = 4.0 \u00d7 \u03c3    (detection sensitivity)\n'
        '    drift     = 0.5 \u00d7 \u03c3    (allowable slack)\n'
        '\n'
        'Algorithm:\n'
        '    S\u207a = 0, S\u207b = 0\n'
        '    \u03bc = average(all values)\n'
        '\n'
        '    For each value x_i:\n'
        '        S\u207a = max(0, S\u207a + (x_i - \u03bc - drift))    // Detect upward shift\n'
        '        S\u207b = max(0, S\u207b + (\u03bc - x_i - drift))    // Detect downward shift\n'
        '\n'
        '        If S\u207a > threshold OR S\u207b > threshold:\n'
        '            \u2192 Declare change point at index i\n'
        '            \u2192 Reset S\u207a = S\u207b = 0')

    add_info_box(doc, 'Why CUSUM? It is sensitive to sustained shifts in mean level, not just individual spikes. '
                      'This makes it ideal for detecting when the pattern of time-travel queries fundamentally '
                      'changes (e.g., users start querying a different historical time period).')

    # 7.3 SmartSnapshotScheduler
    doc.add_heading('7.3 SmartSnapshotScheduler', level=2)
    p = doc.add_paragraph('Decides when to trigger checkpoints based on detected hotspots and change points.')
    add_code_block(doc,
        'Trigger a checkpoint if:\n'
        '    1. Minimum 30 seconds since last snapshot (rate limiting)\n'
        '    AND either:\n'
        '    2a. A change point was detected within the last 5 minutes\n'
        '    2b. A hotspot with density > 1.0 and access_count \u2265 10 exists\n'
        '\n'
        'On trigger:\n'
        '    \u2192 Call CheckpointManager::BeginCheckpoint()\n'
        '    \u2192 Record SNAPSHOT_TRIGGERED in MetricsStore\n'
        '    \u2192 Update scheduled_snapshots list from hotspot centers')

    # 7.4 WALRetentionManager
    doc.add_heading('7.4 WALRetentionManager', level=2)
    p = doc.add_paragraph('Adaptive WAL retention policy based on temporal access patterns.')
    retention_items = [
        ('Hot zone (24 hours): ', 'Full WAL fidelity \u2014 every operation preserved for fast time-travel'),
        ('Cold cutoff (7 days default): ', 'After this threshold, WAL segments can be safely pruned'),
        ('Adaptive extension: ', 'If users query data older than the hot zone, the cold cutoff automatically extends to cover those timestamps plus a configurable buffer'),
    ]
    for prefix, desc_text in retention_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run = p.add_run(desc_text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    # 7.5 TemporalIndexManager
    doc.add_heading('7.5 TemporalIndexManager Orchestrator', level=2)
    add_code_block(doc,
        'class TemporalIndexManager {\n'
        '    void OnTimeTravelQuery(const string& table_name,\n'
        '                           uint64_t target_timestamp,\n'
        '                           const string& db_name);\n'
        '    void PeriodicAnalysis();   // Runs every 30 seconds\n'
        '    string GetSummary() const;\n'
        '    vector<TemporalHotspot> GetCurrentHotspots() const;\n'
        '};', 'cpp')

    doc.add_heading('Periodic Analysis Pipeline (every 30 seconds)', level=4)
    pipeline = [
        'Get all temporal access events from the TemporalAccessTracker',
        'Detect hotspots via DBSCAN clustering on queried timestamps',
        'Generate frequency histogram with 1-minute buckets',
        'Detect change points via CUSUM on histogram values',
        'Update cached hotspots for SHOW AI STATUS',
        'Evaluate snapshot scheduling based on hotspots and change points',
        'Update WAL retention policy based on access patterns',
    ]
    for i, item in enumerate(pipeline):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{i+1}. ')
        run.bold = True
        run.font.color.rgb = PRIMARY
        run = p.add_run(item)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 8. INTEGRATION ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('8. Integration Architecture', level=1)

    p = doc.add_paragraph('This section describes how the AI layer integrates with the existing ChronosDB '
                           'execution engine. The integration is minimal and non-intrusive \u2014 only 9 existing '
                           'files were modified, with approximately 150 lines of changes total.')

    doc.add_heading('8.1 DMLExecutor Integration Pattern', level=2)
    p = doc.add_paragraph('Every DML operation follows this before/after notification pattern:')
    add_code_block(doc,
        '// BEFORE operation\n'
        'ai::DMLEvent ai_event;\n'
        'ai_event.operation = ai::DMLOperation::INSERT;  // or UPDATE, DELETE, SELECT\n'
        'ai_event.table_name = stmt->table_name_;\n'
        'ai_event.start_time_us = now();\n'
        '\n'
        '// Check if Immune System blocks this operation\n'
        'if (!ai::DMLObserverRegistry::Instance().NotifyBefore(ai_event)) {\n'
        '    return ExecutionResult::Error("[IMMUNE] Operation blocked");\n'
        '}\n'
        '\n'
        '// ... execute the actual DML operation ...\n'
        '\n'
        '// AFTER operation\n'
        'ai_event.duration_us = now() - ai_event.start_time_us;\n'
        'ai_event.rows_affected = count;\n'
        'ai::DMLObserverRegistry::Instance().NotifyAfter(ai_event);', 'cpp')

    doc.add_heading('8.2 SELECT-Specific AI Integration', level=2)
    p = doc.add_paragraph('SELECT queries have additional AI integration points beyond the basic observer pattern:')
    add_code_block(doc,
        '// 1. Notify Temporal Index on time-travel queries\n'
        'if (stmt->as_of_timestamp_ > 0) {\n'
        '    ai_mgr.GetTemporalIndexManager()->OnTimeTravelQuery(\n'
        '        table_name, as_of_timestamp, db_name);\n'
        '}\n'
        '\n'
        '// 2. Consult Learning Engine for scan strategy\n'
        'ScanStrategy recommended;\n'
        'if (learning_engine->RecommendScanStrategy(stmt, table_name, recommended)) {\n'
        '    use_index = (recommended == INDEX_SCAN && index_exists);\n'
        '} else {\n'
        '    // Fall back to original heuristic logic\n'
        '}\n'
        '\n'
        '// 3. After SELECT, report feedback with strategy used\n'
        'ai_event.used_index_scan = use_index;\n'
        'ai_event.result_row_count = row_count;\n'
        'DMLObserverRegistry::Instance().NotifyAfter(ai_event);', 'cpp')

    doc.add_heading('8.3 ExecutionEngine Lifecycle Integration', level=2)
    add_code_block(doc,
        '// In ExecutionEngine constructor:\n'
        'ai::AIManager::Instance().Initialize(catalog_, bpm_, log_manager_, cp_mgr);\n'
        '\n'
        '// In ExecutionEngine destructor:\n'
        'ai::AIManager::Instance().Shutdown();', 'cpp')

    p = doc.add_paragraph('This ensures the AI layer starts automatically when the database engine starts '
                           'and shuts down cleanly when the engine stops.')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 9. SQL COMMANDS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('9. SQL Commands', level=1)

    p = doc.add_paragraph('Three new SQL commands were added to expose AI layer status and diagnostics. '
                           'All commands support both English and Franco-Arabic syntax.')

    # SHOW AI STATUS
    doc.add_heading('9.1 SHOW AI STATUS', level=2)
    p = doc.add_paragraph('Displays the status of all three AI subsystems.')
    add_code_block(doc, 'SHOW AI STATUS;\n-- Franco-Arabic: WARENY ZAKA2 7ALAH;', 'sql')

    doc.add_heading('Output Format', level=4)
    add_styled_table(doc,
        ['Component', 'Status', 'Details'],
        [
            ['Learning Engine', 'ACTIVE', '150 queries observed, 2 arms, SeqScan preferred'],
            ['Immune System', 'ACTIVE', '3 anomalies detected, 0 blocked, 0 recoveries'],
            ['Temporal Index', 'ACTIVE', '2 hotspots, 45 time-travel queries, 1 smart snapshot'],
            ['Metrics Recorded', 'INFO', '1,234 events in ring buffer'],
            ['Scheduled Tasks', 'INFO', '3 active background tasks'],
        ])

    # SHOW ANOMALIES
    doc.add_heading('9.2 SHOW ANOMALIES', level=2)
    p = doc.add_paragraph('Shows recent anomaly detections from the Immune System.')
    add_code_block(doc, 'SHOW ANOMALIES;\n-- Franco-Arabic: WARENY SHOZOOZ;', 'sql')

    doc.add_heading('Output Format', level=4)
    add_styled_table(doc,
        ['Table', 'Severity', 'Z-Score', 'Current Rate', 'Mean Rate', 'Detected At'],
        [
            ['users', 'HIGH', '5.23', '150.00/s', '12.50/s', '2026-02-13 14:30:00'],
            ['orders', 'MEDIUM', '3.45', '80.00/s', '20.00/s', '2026-02-13 14:31:00'],
        ])

    # SHOW EXECUTION STATS
    doc.add_heading('9.3 SHOW EXECUTION STATS', level=2)
    p = doc.add_paragraph('Shows the Learning Engine\'s bandit arm statistics and learned preferences.')
    add_code_block(doc, 'SHOW EXECUTION STATS;\n-- Franco-Arabic: WARENY TANFEEZ E7SA2EYAT;', 'sql')

    doc.add_heading('Output Format', level=4)
    add_styled_table(doc,
        ['Strategy', 'Pulls', 'Avg Reward', 'UCB Score'],
        [
            ['Sequential Scan', '85', '0.7234', '0.8901'],
            ['Index Scan', '65', '0.8567', '0.9234'],
            ['Total Queries Observed', '150', '', ''],
        ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 10. UML DIAGRAMS
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('10. UML Diagrams', level=1)

    p = doc.add_paragraph('This section contains PlantUML diagram source code for all major architectural views. '
                           'These diagrams can be rendered using any PlantUML-compatible tool '
                           '(plantuml.com, IDE plugins, or the PlantUML CLI).')

    # Diagram 1: Component Overview
    doc.add_heading('10.1 System Component Diagram', level=2)
    p = doc.add_paragraph('Shows the high-level relationships between ChronosDB core components and '
                           'the AI layer subsystems.')
    add_code_block(doc,
        '@startuml ChronosDB_AI_Components\n'
        '!theme plain\n'
        '\n'
        'package "ChronosDB Core" {\n'
        '    [ExecutionEngine] as EE\n'
        '    [DMLExecutor] as DML\n'
        '    [SystemExecutor] as SYS\n'
        '    [Parser] as PARSER\n'
        '    [BufferPoolManager] as BPM\n'
        '    [LogManager] as LOG\n'
        '    [CheckpointManager] as CP\n'
        '    [TimeTravelEngine] as TT\n'
        '    [Catalog] as CAT\n'
        '}\n'
        '\n'
        'package "AI Layer" {\n'
        '    package "Phase 0: Foundation" {\n'
        '        [AIManager] as AIM\n'
        '        [MetricsStore] as MS\n'
        '        [DMLObserverRegistry] as DOR\n'
        '        [AIScheduler] as SCHED\n'
        '    }\n'
        '    package "Phase 1: Learning Engine" {\n'
        '        [LearningEngine] as LE\n'
        '        [UCB1Bandit] as UCB\n'
        '        [QueryFeatureExtractor] as QFE\n'
        '    }\n'
        '    package "Phase 2: Immune System" {\n'
        '        [ImmuneSystem] as IS\n'
        '        [MutationMonitor] as MM\n'
        '        [UserBehaviorProfiler] as UBP\n'
        '        [AnomalyDetector] as AD\n'
        '        [ResponseEngine] as RE\n'
        '    }\n'
        '    package "Phase 3: Temporal Index" {\n'
        '        [TemporalIndexManager] as TIM\n'
        '        [TemporalAccessTracker] as TAT\n'
        '        [HotspotDetector] as HD\n'
        '        [SmartSnapshotScheduler] as SSS\n'
        '        [WALRetentionManager] as WRM\n'
        '    }\n'
        '}\n'
        '\n'
        'EE --> DML : dispatches DML\n'
        'EE --> SYS : dispatches SHOW\n'
        'EE --> AIM : Initialize/Shutdown\n'
        'DML --> DOR : NotifyBefore/After\n'
        'DML --> LE : RecommendScanStrategy\n'
        'DML --> TIM : OnTimeTravelQuery\n'
        'AIM --> LE : owns\n'
        'AIM --> IS : owns\n'
        'AIM --> TIM : owns\n'
        'LE --> UCB : selects strategy\n'
        'LE --> QFE : extracts features\n'
        'LE --> MS : records metrics\n'
        'LE ..|> DOR : implements IDMLObserver\n'
        'IS --> MM : monitors mutations\n'
        'IS --> UBP : profiles users\n'
        'IS --> AD : detects anomalies\n'
        'IS --> RE : executes responses\n'
        'IS --> MS : records metrics\n'
        'IS ..|> DOR : implements IDMLObserver\n'
        'RE --> TT : auto-recover (HIGH)\n'
        'TIM --> TAT : tracks access\n'
        'TIM --> HD : detects hotspots\n'
        'TIM --> SSS : schedules snapshots\n'
        'TIM --> WRM : manages retention\n'
        'TIM --> MS : records metrics\n'
        'SSS --> CP : BeginCheckpoint\n'
        'SCHED --> IS : periodic analysis\n'
        'SCHED --> TIM : periodic analysis\n'
        'SYS --> AIM : SHOW AI STATUS\n'
        'SYS --> IS : SHOW ANOMALIES\n'
        'SYS --> LE : SHOW EXECUTION STATS\n'
        '@enduml', 'plantuml')

    doc.add_page_break()

    # Diagram 2: Foundation Classes
    doc.add_heading('10.2 Class Diagram \u2014 AI Foundation', level=2)
    add_code_block(doc,
        '@startuml AI_Foundation_Classes\n'
        '\n'
        'abstract class IDMLObserver {\n'
        '    + OnBeforeDML(event: DMLEvent): bool\n'
        '    + OnAfterDML(event: DMLEvent): void\n'
        '}\n'
        '\n'
        'class DMLObserverRegistry {\n'
        '    - observers_: vector<IDMLObserver*>\n'
        '    - mutex_: shared_mutex\n'
        '    + {static} Instance(): DMLObserverRegistry&\n'
        '    + Register(observer): void\n'
        '    + Unregister(observer): void\n'
        '    + NotifyBefore(event): bool\n'
        '    + NotifyAfter(event): void\n'
        '}\n'
        '\n'
        'class MetricsStore {\n'
        '    - buffer_: array<MetricEvent, 8192>\n'
        '    - write_index_: atomic<uint64_t>\n'
        '    + {static} Instance(): MetricsStore&\n'
        '    + Record(event): void\n'
        '    + Query(type, since_us): vector<MetricEvent>\n'
        '}\n'
        '\n'
        'class AIScheduler {\n'
        '    - tasks_: vector<ScheduledTask>\n'
        '    - running_: atomic<bool>\n'
        '    + {static} Instance(): AIScheduler&\n'
        '    + SchedulePeriodic(name, interval, task): TaskId\n'
        '    + Cancel(id): void\n'
        '}\n'
        '\n'
        'class AIManager {\n'
        '    - learning_engine_: unique_ptr<LearningEngine>\n'
        '    - immune_system_: unique_ptr<ImmuneSystem>\n'
        '    - temporal_index_mgr_: unique_ptr<TemporalIndexManager>\n'
        '    + {static} Instance(): AIManager&\n'
        '    + Initialize(...): void\n'
        '    + Shutdown(): void\n'
        '    + GetStatus(): AIStatus\n'
        '}\n'
        '\n'
        'DMLObserverRegistry "1" o-- "*" IDMLObserver\n'
        '@enduml', 'plantuml')

    # Diagram 3: Learning Engine Classes
    doc.add_heading('10.3 Class Diagram \u2014 Learning Engine', level=2)
    add_code_block(doc,
        '@startuml Learning_Engine_Classes\n'
        '\n'
        'abstract class IQueryOptimizer {\n'
        '    + RecommendScanStrategy(stmt, table, out): bool\n'
        '}\n'
        '\n'
        'class LearningEngine {\n'
        '    - catalog_: Catalog*\n'
        '    - feature_extractor_: unique_ptr<QueryFeatureExtractor>\n'
        '    - bandit_: unique_ptr<UCB1Bandit>\n'
        '    - total_queries_: atomic<uint64_t>\n'
        '    + OnAfterDML(event): void\n'
        '    + RecommendScanStrategy(stmt, table, out): bool\n'
        '    + GetSummary(): string\n'
        '    + GetArmStats(): vector<ArmStats>\n'
        '}\n'
        '\n'
        'class UCB1Bandit {\n'
        '    - global_pulls_: atomic<uint64_t>[2]\n'
        '    - global_reward_x10000_: atomic<uint64_t>[2]\n'
        '    - table_stats_: map<string, TableContext>\n'
        '    + SelectStrategy(table): ScanStrategy\n'
        '    + RecordOutcome(strategy, time, table): void\n'
        '    + GetStats(): vector<ArmStats>\n'
        '    - ComputeUCBScore(arm): double\n'
        '    - ComputeReward(time_ms): double\n'
        '}\n'
        '\n'
        'class QueryFeatureExtractor {\n'
        '    + Extract(stmt, table, catalog): QueryFeatures\n'
        '    + EstimateSelectivity(stmt): double\n'
        '}\n'
        '\n'
        'enum ScanStrategy { SEQUENTIAL_SCAN, INDEX_SCAN }\n'
        '\n'
        'LearningEngine --|> IDMLObserver\n'
        'LearningEngine --|> IQueryOptimizer\n'
        'LearningEngine --> QueryFeatureExtractor\n'
        'LearningEngine --> UCB1Bandit\n'
        'UCB1Bandit --> ScanStrategy\n'
        '@enduml', 'plantuml')

    doc.add_page_break()

    # Diagram 4: Immune System Classes
    doc.add_heading('10.4 Class Diagram \u2014 Immune System', level=2)
    add_code_block(doc,
        '@startuml Immune_System_Classes\n'
        '\n'
        'class ImmuneSystem {\n'
        '    - mutation_monitor_: unique_ptr<MutationMonitor>\n'
        '    - user_profiler_: unique_ptr<UserBehaviorProfiler>\n'
        '    - anomaly_detector_: unique_ptr<AnomalyDetector>\n'
        '    - response_engine_: unique_ptr<ResponseEngine>\n'
        '    + OnBeforeDML(event): bool\n'
        '    + OnAfterDML(event): void\n'
        '    + PeriodicAnalysis(): void\n'
        '    + GetRecentAnomalies(max): vector<AnomalyReport>\n'
        '}\n'
        '\n'
        'class MutationMonitor {\n'
        '    - tables_: map<string, deque<MutationEntry>>\n'
        '    + RecordMutation(table, count): void\n'
        '    + GetMutationRate(table): double\n'
        '    + GetHistoricalRates(table): vector<double>\n'
        '}\n'
        '\n'
        'class UserBehaviorProfiler {\n'
        '    - users_: map<string, UserHistory>\n'
        '    + RecordActivity(user, table, is_mutation): void\n'
        '    + GetDeviationScore(user): double\n'
        '}\n'
        '\n'
        'class AnomalyDetector {\n'
        '    + Analyze(monitor): vector<AnomalyReport>\n'
        '    + {static} ComputeZScore(value, history): double\n'
        '}\n'
        '\n'
        'class ResponseEngine {\n'
        '    - blocked_tables_: set<string>\n'
        '    - blocked_users_: set<string>\n'
        '    + Respond(report): void\n'
        '    + IsTableBlocked(table): bool\n'
        '    + IsUserBlocked(user): bool\n'
        '}\n'
        '\n'
        'enum AnomalySeverity { NONE, LOW, MEDIUM, HIGH }\n'
        '\n'
        'ImmuneSystem --|> IDMLObserver\n'
        'ImmuneSystem --> MutationMonitor\n'
        'ImmuneSystem --> UserBehaviorProfiler\n'
        'ImmuneSystem --> AnomalyDetector\n'
        'ImmuneSystem --> ResponseEngine\n'
        '@enduml', 'plantuml')

    # Diagram 5: Temporal Index Classes
    doc.add_heading('10.5 Class Diagram \u2014 Temporal Index Manager', level=2)
    add_code_block(doc,
        '@startuml Temporal_Index_Classes\n'
        '\n'
        'class TemporalIndexManager {\n'
        '    - access_tracker_: unique_ptr<TemporalAccessTracker>\n'
        '    - hotspot_detector_: unique_ptr<HotspotDetector>\n'
        '    - snapshot_scheduler_: unique_ptr<SmartSnapshotScheduler>\n'
        '    - retention_manager_: unique_ptr<WALRetentionManager>\n'
        '    + OnTimeTravelQuery(table, timestamp, db): void\n'
        '    + PeriodicAnalysis(): void\n'
        '    + GetCurrentHotspots(): vector<TemporalHotspot>\n'
        '}\n'
        '\n'
        'class TemporalAccessTracker {\n'
        '    - events_: deque<TemporalAccessEvent>\n'
        '    + RecordAccess(event): void\n'
        '    + GetFrequencyHistogram(bucket_width): vector<FrequencyBucket>\n'
        '    + GetHotTimestamps(k): vector<uint64_t>\n'
        '}\n'
        '\n'
        'class HotspotDetector {\n'
        '    + DetectHotspots(events): vector<TemporalHotspot>\n'
        '    + DetectChangePoints(rates, timestamps): vector<uint64_t>\n'
        '}\n'
        '\n'
        'class SmartSnapshotScheduler {\n'
        '    - checkpoint_mgr_: CheckpointManager*\n'
        '    - last_snapshot_time_us_: uint64_t\n'
        '    + Evaluate(hotspots, change_points): void\n'
        '    + GetScheduledSnapshots(): vector<uint64_t>\n'
        '}\n'
        '\n'
        'class WALRetentionManager {\n'
        '    - log_manager_: LogManager*\n'
        '    + ComputePolicy(tracker): RetentionPolicy\n'
        '    + UpdatePolicy(policy): void\n'
        '}\n'
        '\n'
        'TemporalIndexManager --> TemporalAccessTracker\n'
        'TemporalIndexManager --> HotspotDetector\n'
        'TemporalIndexManager --> SmartSnapshotScheduler\n'
        'TemporalIndexManager --> WALRetentionManager\n'
        '@enduml', 'plantuml')

    doc.add_page_break()

    # Diagram 6: SELECT Sequence
    doc.add_heading('10.6 Sequence Diagram \u2014 SELECT with AI', level=2)
    p = doc.add_paragraph('Shows the complete flow of a SELECT query through the AI layer, including '
                           'strategy recommendation and feedback recording.')
    add_code_block(doc,
        '@startuml SELECT_AI_Sequence\n'
        '\n'
        'actor User\n'
        'participant DMLExecutor\n'
        'participant LearningEngine\n'
        'participant UCB1Bandit\n'
        'participant DMLObserverRegistry\n'
        'participant MetricsStore\n'
        '\n'
        'User -> DMLExecutor: SELECT * FROM users WHERE id = 5;\n'
        '\n'
        'DMLExecutor -> LearningEngine: RecommendScanStrategy(stmt, "users")\n'
        'LearningEngine -> UCB1Bandit: SelectStrategy("users")\n'
        'UCB1Bandit --> LearningEngine: INDEX_SCAN\n'
        'LearningEngine --> DMLExecutor: true, INDEX_SCAN\n'
        '\n'
        'DMLExecutor -> DMLExecutor: Execute with IndexScanExecutor\n'
        '\n'
        'DMLExecutor -> DMLObserverRegistry: NotifyAfter(event)\n'
        'DMLObserverRegistry -> LearningEngine: OnAfterDML(event)\n'
        'LearningEngine -> UCB1Bandit: RecordOutcome(INDEX_SCAN, 5ms, "users")\n'
        'UCB1Bandit -> UCB1Bandit: Update reward: 1/(1+5/100) = 0.95\n'
        'LearningEngine -> MetricsStore: Record(SELECT metric)\n'
        'DMLObserverRegistry --> DMLExecutor: done\n'
        '\n'
        'DMLExecutor --> User: ResultSet (1 row)\n'
        '@enduml', 'plantuml')

    # Diagram 7: Anomaly Sequence
    doc.add_heading('10.7 Sequence Diagram \u2014 Anomaly Detection & Auto-Recovery', level=2)
    p = doc.add_paragraph('Shows the complete anomaly detection and auto-recovery flow when a mass DELETE '
                           'triggers HIGH severity.')
    add_code_block(doc,
        '@startuml Anomaly_Recovery_Sequence\n'
        '\n'
        'actor Attacker\n'
        'participant DMLExecutor\n'
        'participant DMLObserverRegistry\n'
        'participant ImmuneSystem\n'
        'participant MutationMonitor\n'
        'participant AnomalyDetector\n'
        'participant ResponseEngine\n'
        'participant TimeTravelEngine\n'
        '\n'
        '== Phase 1: Malicious Mass Delete ==\n'
        'Attacker -> DMLExecutor: DELETE FROM users WHERE id > 0;\n'
        'DMLExecutor -> DMLObserverRegistry: NotifyBefore(DELETE)\n'
        'DMLObserverRegistry -> ImmuneSystem: OnBeforeDML(event)\n'
        'ImmuneSystem --> DMLObserverRegistry: true (allowed)\n'
        'DMLExecutor -> DMLExecutor: Execute DELETE (500 rows)\n'
        'DMLExecutor -> DMLObserverRegistry: NotifyAfter(DELETE, 500 rows)\n'
        'DMLObserverRegistry -> ImmuneSystem: OnAfterDML(event)\n'
        'ImmuneSystem -> MutationMonitor: RecordMutation("users", 500)\n'
        '\n'
        '== Phase 2: Periodic Analysis Detects Anomaly ==\n'
        'ImmuneSystem -> AnomalyDetector: Analyze(mutation_monitor)\n'
        'AnomalyDetector -> MutationMonitor: GetMutationRate("users") = 500/s\n'
        'AnomalyDetector -> MutationMonitor: GetHistoricalRates("users")\n'
        'AnomalyDetector -> AnomalyDetector: Z-score = (500 - 2) / 0.8 = 622.5\n'
        'AnomalyDetector --> ImmuneSystem: AnomalyReport(HIGH, z=622.5)\n'
        '\n'
        'ImmuneSystem -> ResponseEngine: Respond(report)\n'
        'ResponseEngine -> ResponseEngine: BlockTable("users")\n'
        'ResponseEngine -> TimeTravelEngine: RecoverTo(now - 60s, db)\n'
        'TimeTravelEngine --> ResponseEngine: success\n'
        'ResponseEngine -> ResponseEngine: UnblockTable("users")\n'
        '\n'
        '== Phase 3: Subsequent Attacks Blocked ==\n'
        'Attacker -> DMLExecutor: DELETE FROM users WHERE id > 0;\n'
        'DMLExecutor -> DMLObserverRegistry: NotifyBefore(DELETE)\n'
        'DMLObserverRegistry -> ImmuneSystem: OnBeforeDML(event)\n'
        'ImmuneSystem -> ResponseEngine: IsTableBlocked("users")\n'
        'ResponseEngine --> ImmuneSystem: true\n'
        'ImmuneSystem --> DMLObserverRegistry: false (BLOCKED)\n'
        'DMLExecutor --> Attacker: ERROR: [IMMUNE] DELETE blocked\n'
        '@enduml', 'plantuml')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 11. FILE INVENTORY
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('11. File Inventory', level=1)

    p = doc.add_paragraph('The AI layer consists of 35 new files and modifications to 9 existing files. '
                           'All new files are organized under the src/ai/ and src/include/ai/ directories.')

    doc.add_heading('Phase 0: Shared Foundation (9 files)', level=2)
    add_styled_table(doc,
        ['File', 'Lines', 'Purpose'],
        [
            ['src/include/ai/ai_config.h', '~60', 'All AI constants centralized'],
            ['src/include/ai/metrics_store.h', '~80', 'MetricsStore header'],
            ['src/ai/metrics_store.cpp', '~110', 'MetricsStore ring buffer implementation'],
            ['src/include/ai/dml_observer.h', '~90', 'Observer interface + registry header'],
            ['src/ai/dml_observer.cpp', '~55', 'Observer registry implementation'],
            ['src/include/ai/ai_scheduler.h', '~65', 'Background task scheduler header'],
            ['src/ai/ai_scheduler.cpp', '~100', 'Scheduler implementation'],
            ['src/include/ai/ai_manager.h', '~75', 'Top-level coordinator header'],
            ['src/ai/ai_manager.cpp', '~100', 'Manager implementation'],
        ])

    doc.add_heading('Phase 1: Learning Engine (6 files)', level=2)
    add_styled_table(doc,
        ['File', 'Lines', 'Purpose'],
        [
            ['src/include/ai/learning/query_features.h', '~50', 'Feature extractor header'],
            ['src/ai/learning/query_features.cpp', '~80', 'Feature extraction implementation'],
            ['src/include/ai/learning/bandit.h', '~70', 'UCB1 Bandit header'],
            ['src/ai/learning/bandit.cpp', '~130', 'Bandit implementation with per-table context'],
            ['src/include/ai/learning/learning_engine.h', '~80', 'Learning orchestrator header'],
            ['src/ai/learning/learning_engine.cpp', '~90', 'Learning implementation'],
        ])

    doc.add_heading('Phase 2: Immune System (10 files)', level=2)
    add_styled_table(doc,
        ['File', 'Lines', 'Purpose'],
        [
            ['src/include/ai/immune/mutation_monitor.h', '~60', 'Mutation tracking header'],
            ['src/ai/immune/mutation_monitor.cpp', '~100', 'Rolling window monitor implementation'],
            ['src/include/ai/immune/user_profiler.h', '~60', 'User behavior profiling header'],
            ['src/ai/immune/user_profiler.cpp', '~140', 'Profiler with deviation scoring'],
            ['src/include/ai/immune/anomaly_detector.h', '~55', 'Z-score anomaly detection header'],
            ['src/ai/immune/anomaly_detector.cpp', '~80', 'Detector implementation'],
            ['src/include/ai/immune/response_engine.h', '~60', 'Graduated response header'],
            ['src/ai/immune/response_engine.cpp', '~130', 'Response engine with auto-recovery'],
            ['src/include/ai/immune/immune_system.h', '~65', 'Immune orchestrator header'],
            ['src/ai/immune/immune_system.cpp', '~120', 'Immune system implementation'],
        ])

    doc.add_heading('Phase 3: Temporal Index (10 files)', level=2)
    add_styled_table(doc,
        ['File', 'Lines', 'Purpose'],
        [
            ['src/include/ai/temporal/access_tracker.h', '~55', 'Access tracking header'],
            ['src/ai/temporal/access_tracker.cpp', '~105', 'Tracker with frequency analysis'],
            ['src/include/ai/temporal/hotspot_detector.h', '~55', 'Hotspot detection header'],
            ['src/ai/temporal/hotspot_detector.cpp', '~170', 'DBSCAN + CUSUM implementation'],
            ['src/include/ai/temporal/snapshot_scheduler.h', '~50', 'Snapshot scheduling header'],
            ['src/ai/temporal/snapshot_scheduler.cpp', '~100', 'Smart scheduler implementation'],
            ['src/include/ai/temporal/retention_manager.h', '~50', 'WAL retention header'],
            ['src/ai/temporal/retention_manager.cpp', '~65', 'Adaptive retention implementation'],
            ['src/include/ai/temporal/temporal_index_manager.h', '~70', 'Temporal orchestrator header'],
            ['src/ai/temporal/temporal_index_manager.cpp', '~140', 'Manager implementation'],
        ])

    doc.add_heading('Modified Existing Files (9 files)', level=2)
    add_styled_table(doc,
        ['File', 'Changes', 'Purpose'],
        [
            ['src/include/parser/token.h', '+4 tokens', 'AI, ANOMALIES, EXECUTION, STATS'],
            ['src/parser/lexer.cpp', '+12 keywords', 'Franco-Arabic + English keyword mappings'],
            ['src/include/parser/statement.h', '+3 types, +3 classes', 'SHOW_AI_STATUS, SHOW_ANOMALIES, SHOW_EXECUTION_STATS'],
            ['src/parser/parser.cpp', '+18 lines', 'SHOW parsing for new AI commands'],
            ['src/execution/execution_engine.cpp', '+15 lines', 'AI init/shutdown + 3 dispatch entries'],
            ['src/include/execution/system_executor.h', '+6 lines', '3 new SHOW handler declarations'],
            ['src/execution/executors/system_executor.cpp', '+100 lines', 'SHOW AI STATUS/ANOMALIES/STATS implementations'],
            ['src/execution/executors/dml_executor.cpp', '+80 lines', 'Observer hooks + AI consultation in all DML methods'],
            ['src/include/common/logger.h', '+1 line', 'Added missing <iostream> include'],
        ])

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 12. ALGORITHM REFERENCE
    # ═══════════════════════════════════════════════════════════════
    doc.add_heading('12. Algorithm Reference', level=1)

    p = doc.add_paragraph('This section provides a consolidated mathematical reference for all algorithms '
                           'used in the AI layer.')

    doc.add_heading('12.1 UCB1 Multi-Armed Bandit', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    run = p.add_run('Choose between Sequential Scan and Index Scan for each query without knowing '
                     'in advance which strategy is faster for the given workload.')

    p = doc.add_paragraph()
    run = p.add_run('Solution: ')
    run.bold = True
    run = p.add_run('Treat each scan strategy as an "arm" of a slot machine. Use UCB1 to balance '
                     'exploitation (use the best-known strategy) with exploration (try uncertain strategies '
                     'to gather more data).')

    add_code_block(doc,
        'Selection:  a* = argmax_a [ Q\u0302(a) + c \u00b7 \u221a(ln N / N_a) ]\n'
        '\n'
        '    Q\u0302(a) = total_reward(a) / N_a          (estimated arm value)\n'
        '    c    = \u221a2 \u2248 1.414                      (exploration constant)\n'
        '    N    = total pulls across all arms      (global experience)\n'
        '    N_a  = pulls for arm a                  (arm-specific experience)\n'
        '\n'
        'Reward:  r(t) = 1.0 / (1.0 + t_ms / 100.0)\n'
        '\n'
        '    10ms   \u2192 0.91 (excellent)\n'
        '    100ms  \u2192 0.50 (average)\n'
        '    1000ms \u2192 0.09 (poor)\n'
        '\n'
        'Properties:\n'
        '    \u2022 Bounded in (0, 1]\n'
        '    \u2022 Smooth and differentiable\n'
        '    \u2022 Per-table contextual learning after 10+ pulls per arm\n'
        '    \u2022 Exploration phase: first 30 queries use existing heuristics')

    doc.add_heading('12.2 Z-Score Anomaly Detection', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    run = p.add_run('Detect when a table\'s mutation rate is abnormally high, which may indicate '
                     'an attack, a bug, or an accidental mass operation.')

    add_code_block(doc,
        'z = (x - \u03bc) / \u03c3\n'
        '\n'
        '    x = current mutation rate (mutations/sec)\n'
        '    \u03bc = mean of last 100 intervals\n'
        '    \u03c3 = standard deviation of last 100 intervals\n'
        '\n'
        'Statistical Interpretation:\n'
        '    z < 2.0  \u2192 NONE   (within 95.4% of normal variation)\n'
        '    z \u2265 2.0  \u2192 LOW    (4.6% false positive rate - log warning)\n'
        '    z \u2265 3.0  \u2192 MEDIUM (0.3% false positive - block mutations)\n'
        '    z \u2265 4.0  \u2192 HIGH   (0.006% false positive - auto-recover)')

    doc.add_heading('12.3 DBSCAN Clustering (1D Simplified)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    run = p.add_run('Find temporal hotspots \u2014 time ranges that are frequently queried in time-travel operations.')

    add_code_block(doc,
        'Simplified 1D DBSCAN:\n'
        '\n'
        '    Parameters: epsilon = 60s, min_points = 5\n'
        '\n'
        '    1. Sort queried timestamps\n'
        '    2. Walk through sorted list\n'
        '    3. Group consecutive timestamps within epsilon distance\n'
        '    4. Clusters with \u2265 min_points become hotspots\n'
        '    5. For each hotspot: center = mean, density = count / range\n'
        '\n'
        '    Complexity: O(n) for sorted 1D data\n'
        '    (vs. O(n\u00b2) for general multi-dimensional DBSCAN)')

    doc.add_heading('12.4 CUSUM Change-Point Detection', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Problem: ')
    run.bold = True
    run = p.add_run('Detect when the pattern of temporal queries fundamentally shifts, indicating '
                     'a good moment to create a new snapshot.')

    add_code_block(doc,
        'CUSUM Algorithm:\n'
        '\n'
        '    Parameters:\n'
        '        threshold = 4.0\u03c3    (detection sensitivity)\n'
        '        drift     = 0.5\u03c3    (allowable slack)\n'
        '\n'
        '    Initialize: S\u207a = 0, S\u207b = 0\n'
        '\n'
        '    For each observation x_i:\n'
        '        S\u207a = max(0, S\u207a + (x_i - \u03bc - drift))    // upward shift\n'
        '        S\u207b = max(0, S\u207b + (\u03bc - x_i - drift))    // downward shift\n'
        '\n'
        '        If S\u207a > threshold OR S\u207b > threshold:\n'
        '            Declare change point\n'
        '            Reset S\u207a = S\u207b = 0\n'
        '\n'
        '    Why CUSUM:\n'
        '        \u2022 Sensitive to sustained mean shifts, not individual spikes\n'
        '        \u2022 Low false positive rate with 4\u03c3 threshold\n'
        '        \u2022 Computationally efficient: O(n) single pass')

    # ═══════════════════════════════════════════════════════════════
    # FOOTER
    # ═══════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_horizontal_rule(doc)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(20)
    run = footer.add_run('ChronosDB AI Layer Technical Documentation')
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_TEXT
    run.italic = True

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver.add_run('Version 1.0  \u2022  February 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_TEXT
    run.italic = True

    built = doc.add_paragraph()
    built.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = built.add_run('Built with C++20  \u2022  ~4,200 lines of new AI code  \u2022  35 new files')
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_TEXT
    run.italic = True

    return doc


if __name__ == '__main__':
    print('Generating ChronosDB AI Layer Documentation (.docx)...')
    doc = build_document()

    output_path = os.path.join(os.path.dirname(__file__), 'ChronosDB_AI_Layer_Documentation.docx')
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    print('Done!')
