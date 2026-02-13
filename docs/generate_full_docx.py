"""
Generate professionally styled Word document for ChronosDB Full System Documentation.
Parses the Markdown file and converts it to a styled .docx.
Includes PlantUML diagram rendering via the PlantUML web service.
"""

import os
import re
import zlib
import tempfile
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

try:
    from urllib.request import urlopen, Request
    from urllib.error import URLError
    HAS_URLLIB = True
except ImportError:
    HAS_URLLIB = False

# ── Color Palette ──────────────────────────────────────────────────
PRIMARY      = RGBColor(0x1A, 0x56, 0xDB)
ACCENT       = RGBColor(0x0E, 0x9F, 0x6E)
DARK         = RGBColor(0x1F, 0x29, 0x37)
MEDIUM       = RGBColor(0x4B, 0x55, 0x63)
LIGHT_TEXT   = RGBColor(0x6B, 0x72, 0x80)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)


# ── PlantUML Rendering ────────────────────────────────────────────
def _encode6bit(b):
    if b < 10: return chr(48 + b)
    b -= 10
    if b < 26: return chr(65 + b)
    b -= 26
    if b < 26: return chr(97 + b)
    b -= 26
    if b == 0: return '-'
    if b == 1: return '_'
    return '?'

def plantuml_encode(text):
    """Encode PlantUML text for the web service URL."""
    compressed = zlib.compress(text.encode('utf-8'))
    compressed = compressed[2:-4]  # Strip zlib header/checksum
    result = ''
    i = 0
    while i < len(compressed):
        if i + 2 < len(compressed):
            b1, b2, b3 = compressed[i], compressed[i+1], compressed[i+2]
            result += _encode6bit(b1 >> 2)
            result += _encode6bit(((b1 & 0x3) << 4) | (b2 >> 4))
            result += _encode6bit(((b2 & 0xF) << 2) | (b3 >> 6))
            result += _encode6bit(b3 & 0x3F)
        elif i + 1 < len(compressed):
            b1, b2 = compressed[i], compressed[i+1]
            result += _encode6bit(b1 >> 2)
            result += _encode6bit(((b1 & 0x3) << 4) | (b2 >> 4))
            result += _encode6bit((b2 & 0xF) << 2)
        else:
            b1 = compressed[i]
            result += _encode6bit(b1 >> 2)
            result += _encode6bit((b1 & 0x3) << 4)
        i += 3
    return result

def render_plantuml(text, output_dir, retries=3):
    """Render PlantUML text to a PNG image file. Returns path or None."""
    if not HAS_URLLIB:
        return None
    import time
    for attempt in range(retries):
        try:
            if attempt > 0:
                time.sleep(2 * attempt)
            encoded = plantuml_encode(text)
            url = f'http://www.plantuml.com/plantuml/png/{encoded}'
            req = Request(url, headers={'User-Agent': 'ChronosDB-DocGen/1.0'})
            response = urlopen(req, timeout=30)
            img_data = response.read()
            if len(img_data) < 100:
                continue
            fd, path = tempfile.mkstemp(suffix='.png', dir=output_dir)
            with os.fdopen(fd, 'wb') as f:
                f.write(img_data)
            return path
        except Exception as e:
            print(f'  [WARN] PlantUML render attempt {attempt + 1} failed: {e}')
            if attempt == retries - 1:
                return None
    return None

def add_plantuml_diagram(doc, plantuml_text, output_dir):
    """Render and add a PlantUML diagram as an image to the document."""
    img_path = render_plantuml(plantuml_text, output_dir)
    if img_path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(6.0))
            print(f'  [OK] Embedded PlantUML diagram as image')
            try:
                os.remove(img_path)
            except:
                pass
            return True
        except Exception as e:
            print(f'  [WARN] Failed to embed image: {e}')
            try:
                os.remove(img_path)
            except:
                pass
    # Fallback: add as styled code block with diagram label
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run('\u25bc UML Diagram (PlantUML Source) \u25bc')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    add_code_block(doc, plantuml_text)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run('Render at: plantuml.com/plantuml')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = LIGHT_TEXT
    return False


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color, before, after) in {
        1: (24, PRIMARY, 24, 12),
        2: (18, DARK, 18, 8),
        3: (14, ACCENT, 14, 6),
        4: (12, MEDIUM, 10, 4),
    }.items():
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.keep_with_next = True


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        set_cell_shading(cell, '1A56DB')

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            if c_idx >= len(table.columns):
                continue
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4FF')

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = DARK
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F3F4F6"/>')
    pPr.append(shading)


def add_info_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK
    run.italic = True
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="E8F4FD"/>')
    pPr.append(shading)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="4" w:color="1A56DB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_paragraph_with_formatting(doc, text):
    """Add a paragraph, handling **bold** and `code` inline formatting."""
    p = doc.add_paragraph()
    # Split by **bold** and `code` patterns
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x9B, 0x17, 0x4D)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    return p


def add_bullet_with_formatting(doc, text, level=0):
    """Add a bullet point, handling **bold** and `code` inline formatting."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
    p.paragraph_format.space_after = Pt(3)
    # Clear default runs
    for r in p.runs:
        r.text = ''

    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x9B, 0x17, 0x4D)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
    return p


def parse_markdown_to_docx(md_path, doc):
    """Parse Markdown file and add content to docx."""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    output_dir = os.path.dirname(os.path.abspath(md_path))
    i = 0
    in_code_block = False
    is_plantuml = False
    code_buffer = []
    in_table = False
    table_headers = []
    table_rows = []
    diagram_count = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                if is_plantuml:
                    diagram_count += 1
                    if diagram_count > 1:
                        import time; time.sleep(1)  # Avoid rate limiting
                    print(f'  Rendering PlantUML diagram {diagram_count}...')
                    add_plantuml_diagram(doc, '\n'.join(code_buffer), output_dir)
                else:
                    add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
                is_plantuml = False
            else:
                # Flush any pending table
                if in_table:
                    add_styled_table(doc, table_headers, table_rows)
                    in_table = False
                    table_headers = []
                    table_rows = []
                in_code_block = True
                is_plantuml = line.strip().lower().startswith('```plantuml')
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Table rows
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Check if separator row (---|---)
            if all(re.match(r'^[-:]+$', c) for c in cells):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                add_styled_table(doc, table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []

        # Headings
        if line.startswith('# ') and not line.startswith('##'):
            # Skip the main title — we have a cover page
            i += 1
            continue
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            # Remove markdown link syntax from heading
            heading_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', heading_text)
            doc.add_heading(heading_text, level=1)
            i += 1
            continue
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        elif line.startswith('#### '):
            heading_text = line[5:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # Bullet points
        if line.startswith('- '):
            add_bullet_with_formatting(doc, line[2:])
            i += 1
            continue

        # Numbered lists
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            num = m.group(1)
            text = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f'{num}. ')
            run.bold = True
            run.font.color.rgb = PRIMARY
            run.font.size = Pt(11)
            # Handle formatting in rest
            parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x9B, 0x17, 0x4D)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(11)
                    run.font.name = 'Calibri'
            i += 1
            continue

        # Empty lines
        if line.strip() == '':
            i += 1
            continue

        # Italic info text (starts with *)
        if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            text = line.strip('*').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = LIGHT_TEXT
            i += 1
            continue

        # Regular paragraph
        add_paragraph_with_formatting(doc, line)
        i += 1

    # Flush remaining table
    if in_table:
        add_styled_table(doc, table_headers, table_rows)


def build_document(md_path):
    doc = Document()
    style_document(doc)

    # ═══ COVER PAGE ═══
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ChronosDB')
    run.font.size = Pt(42)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.name = 'Calibri'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Complete System Technical Documentation')
    run.font.size = Pt(20)
    run.font.color.rgb = MEDIUM
    run.font.name = 'Calibri'

    doc.add_paragraph()

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run('\u2501' * 40)
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(14)

    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        'Storage  \u2022  Buffer Pool  \u2022  Parser  \u2022  Execution Engine  \u2022  Recovery\n'
        'Time Travel  \u2022  Network  \u2022  RBAC  \u2022  AI Layer  \u2022  Concurrency'
    )
    run.font.size = Pt(12)
    run.font.color.rgb = LIGHT_TEXT
    run.font.name = 'Calibri'

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Version 1.0  \u2022  February 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = LIGHT_TEXT

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta2.add_run('C++20  \u2022  CMake + Ninja  \u2022  Windows / Linux')
    run.font.size = Pt(11)
    run.font.color.rgb = LIGHT_TEXT

    doc.add_page_break()

    # ═══ TABLE OF CONTENTS ═══
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Project Overview'),
        ('2.', 'System Architecture'),
        ('3.', 'Storage Layer'),
        ('4.', 'Buffer Pool Management'),
        ('5.', 'Catalog & Metadata'),
        ('6.', 'Parser & Lexer'),
        ('7.', 'Execution Engine'),
        ('8.', 'Concurrency Control'),
        ('9.', 'Recovery & Write-Ahead Logging'),
        ('10.', 'Time Travel Engine'),
        ('11.', 'Network Layer'),
        ('12.', 'Authentication & RBAC'),
        ('13.', 'AI Layer \u2014 Self-Learning Execution Engine'),
        ('14.', 'AI Layer \u2014 Immune System'),
        ('15.', 'AI Layer \u2014 Temporal Index Manager'),
        ('16.', 'AI Shared Foundation'),
        ('17.', 'Command-Line Interfaces'),
        ('18.', 'SQL Commands Reference'),
        ('19.', 'UML Diagrams'),
        ('20.', 'File Inventory'),
        ('21.', 'Algorithm Reference'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f'{num}  ')
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = PRIMARY
        run.font.name = 'Calibri'
        run = p.add_run(title_text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        run.font.name = 'Calibri'

    doc.add_page_break()

    # ═══ MAIN CONTENT (parsed from Markdown) ═══
    parse_markdown_to_docx(md_path, doc)

    # ═══ FOOTER ═══
    doc.add_page_break()
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    for text in [
        'ChronosDB Complete System Technical Documentation',
        'Version 1.0  \u2022  February 2026',
        'C++20  \u2022  100+ source files  \u2022  15,000+ lines of code',
        '21 chapters  \u2022  9 UML diagrams  \u2022  8 algorithms documented',
    ]:
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = LIGHT_TEXT
        run.italic = True

    return doc


if __name__ == '__main__':
    print('Generating ChronosDB Full System Documentation (.docx)...')

    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, 'ChronosDB_Full_Documentation.md')
    output_path = os.path.join(script_dir, 'ChronosDB_Full_Documentation.docx')

    doc = build_document(md_path)
    doc.save(output_path)

    file_size = os.path.getsize(output_path)
    print(f'Document saved to: {output_path}')
    print(f'File size: {file_size:,} bytes ({file_size/1024:.0f} KB)')
    print('Done!')
